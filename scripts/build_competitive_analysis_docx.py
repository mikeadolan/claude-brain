#!/usr/bin/env python3
"""Build Claude Brain Competitive Analysis DOCX."""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

OUTPUT = os.path.expanduser(
    "~/Dropbox/Documents/AI/Claude/claude-brain/exports/Claude_Brain_Competitive_Analysis.docx"
)

doc = Document()
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.space_before = Pt(0)

# --- Helper functions ---
def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    return h

def add_para(text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    return p

def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
    return table


# ============================================================
# TITLE PAGE
# ============================================================
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_heading('Claude Brain', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

subtitle = doc.add_heading('Competitive Analysis & Strategic Assessment', level=1)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in subtitle.runs:
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x44, 0x44, 0x66)

doc.add_paragraph()
date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = date_p.add_run('March 9, 2026')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph()
conf = doc.add_paragraph()
conf.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = conf.add_run('CONFIDENTIAL')
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

doc.add_page_break()

# ============================================================
# EXECUTIVE SUMMARY
# ============================================================
add_heading('Executive Summary')

add_para(
    'This document presents a comprehensive competitive analysis of the AI memory '
    'tool landscape as of March 2026, positioning claude-brain against all known '
    'competitors across three categories: general AI memory layers, Claude Code-specific '
    'memory tools, and built-in memory features of AI coding assistants.'
)

add_para(
    'The AI memory space has seen rapid growth: nine major open-source projects, '
    'over $45M in venture capital funding, and every AI coding tool adding some form '
    'of persistent memory. Despite this crowded landscape, claude-brain occupies a '
    'genuinely unique position.'
)

add_heading('Key Findings', level=2)

add_bullet(
    'It is a general-purpose AI assistant platform (WhatsApp, Telegram, Slack), '
    'not a Claude Code memory tool. Different product category entirely.',
    bold_prefix='OpenClaw (287K stars) is not a competitor. '
)
add_bullet(
    'claude-mem (33.5K stars) captures tool observations only, not full transcripts. '
    'No structured knowledge, no proactive features, no governance.',
    bold_prefix='The real Claude Code competitor is claude-mem. '
)
add_bullet(
    'Mem0 ($24M funded, 49K stars) is an API for developers building AI apps, '
    'not a personal development memory system.',
    bold_prefix='The VC-funded heavyweights serve a different market. '
)
add_bullet(
    'Full transcript capture + structured governance + semantic search + proactive '
    'email digests + cross-project isolation + local-first + zero cost. '
    'No single competitor offers this combination.',
    bold_prefix='claude-brain\'s combination is unique. '
)
add_bullet(
    'Anthropic acquired Bun (~$100-200M) because Claude Code depends on it. '
    'Becoming the memory layer Claude Code power users depend on is a viable path.',
    bold_prefix='The acquisition precedent is real but requires scale. '
)

doc.add_page_break()

# ============================================================
# MARKET LANDSCAPE
# ============================================================
add_heading('Market Landscape')

add_heading('Category 1: General AI Memory Layers', level=2)
add_para(
    'These are VC-funded platforms providing memory-as-a-service APIs for developers '
    'building AI applications. They serve a fundamentally different audience than '
    'claude-brain (app developers vs. individual developers using Claude Code).'
)

add_table(
    ['Project', 'Stars', 'Funding', 'Architecture', 'Pricing'],
    [
        ['Mem0', '~49K', '$24M Series A', 'Vector DB + Neo4j graph + KV store', 'Free / $19mo / $249mo'],
        ['Khoj', '~33K', '$500K (YC)', 'RAG over personal files', 'Free / cloud tiers'],
        ['Letta (MemGPT)', '~21K', '$10M seed', 'Stateful agent framework', 'Open source / Cloud'],
        ['Zep / Graphiti', '~20K', '$500K+ (YC)', 'Temporal knowledge graph', 'Graphiti OSS / Zep Cloud'],
        ['Cognee', '~12.7K', '$9M seed', 'Vector + graph + cognitive', 'Open source / Enterprise'],
        ['Supermemory', '~10K', '$2.6M seed', 'Memory engine + RAG', 'OSS + cloud API'],
        ['MemOS', '~5.2K', 'Unknown', 'Memory OS abstraction', 'OSS + cloud'],
        ['OpenMemory', '~3-5K', 'Community', 'Cognitive engine, MCP-native', 'Free, self-hosted'],
        ['LangMem', '~1.2K', 'LangChain backed', 'SDK for LangGraph', 'Free, open source'],
    ]
)

doc.add_paragraph()
add_para(
    'Key takeaway: Mem0 is the category leader at 49K stars and $24M funding. '
    'However, all of these are general-purpose memory APIs for developers building '
    'AI applications, not personal memory systems for individual developers using Claude Code.',
    italic=True
)

doc.add_paragraph()

add_heading('Category 2: Claude Code Memory Tools (Direct Competitors)', level=2)
add_para(
    'These tools specifically add persistent memory to Claude Code. This is claude-brain\'s '
    'direct competitive space.'
)

add_table(
    ['Tool', 'Stars', 'Architecture', 'Key Differentiator'],
    [
        ['Claude Code built-in', 'N/A', 'MEMORY.md (200-line cap)', 'Native, zero setup. No semantic search.'],
        ['claude-mem', '~33.5K', 'SQLite + Chroma + Bun worker', 'Most popular plugin. Heavy setup.'],
        ['memory-mcp', '--', 'Two-tier JSON + auto CLAUDE.md', 'Haiku LLM extracts. ~$0.05/day.'],
        ['mcp-memory-service', '~1.3K', 'SQLite-vec + knowledge graph', 'Enterprise, 20+ client support.'],
        ['memsearch (Zilliz)', '--', 'Markdown + Milvus vector DB', 'Extracted from OpenClaw. Milvus co.'],
        ['context-manager', '--', 'Compaction hooks + checkpoints', 'Prevents context loss, not memory.'],
        ['claude-brain', '0 (private)', 'SQLite + FTS5 + numpy + MCP', 'Full governance: transcripts, decisions, facts, digests.'],
    ]
)

doc.add_paragraph()
add_para(
    'Key takeaway: claude-mem at 33.5K stars is the incumbent. However, it captures '
    'tool observations only (not full transcripts), has no structured knowledge tables, '
    'no proactive features, and no governance system.',
    italic=True
)

doc.add_paragraph()

add_heading('Category 3: Built-in Memory in AI Coding Tools', level=2)

add_table(
    ['Tool', 'Memory Approach', 'Key Limitation'],
    [
        ['Cursor', '.mdc rules + Generate Memories', 'Flat markdown, no semantic search'],
        ['Windsurf', 'Auto-memories + AGENTS.md rules', 'Local-only, quality varies'],
        ['GitHub Copilot', 'Repo-scoped facts, 28-day expiry', 'Aggressive expiry, no user control'],
        ['Cline', 'Memory Bank methodology (markdown)', 'Convention, not a feature'],
        ['Aider', 'Conventions file + repo map', 'Zero persistent cross-session memory'],
        ['Kiro (AWS)', 'Specs ARE the memory', 'Spec-driven only, not general'],
        ['Continue.dev', 'MCP integration, no built-in', 'Must bring your own solution'],
    ]
)

doc.add_paragraph()
add_para(
    'Key takeaway: Every tool has bolted on some form of memory, but they are all '
    'shallow. Copilot\'s 28-day expiry is particularly telling. Nobody has solved this well.',
    italic=True
)

doc.add_page_break()

# ============================================================
# OPENCLAW CORRECTION
# ============================================================
add_heading('OpenClaw: Critical Correction')

add_para(
    'Our initial strategy was framed as an "OpenClaw killer." This was based on a '
    'misunderstanding of what OpenClaw actually is.'
)

add_heading('What OpenClaw Actually Is', level=2)
add_bullet('General-purpose AI assistant platform (WhatsApp, Telegram, Slack, Discord, iMessage)')
add_bullet('287K GitHub stars (most-starred software project, surpassing React)')
add_bullet('Model-agnostic orchestration (Claude, GPT, DeepSeek, Gemini, Ollama)')
add_bullet('Built by Peter Steinberger (PSPDFKit founder), acqui-hired by OpenAI Feb 2026')
add_bullet('Memory is its weakest component, called "broken by default" by users')
add_bullet('512 security vulnerabilities (8 critical), plain-text credentials')
add_bullet('135,000+ instances exposed on the public internet')
add_bullet('$10-30/month infrastructure cost')

add_para('')
add_para(
    'OpenClaw is not a Claude Code memory tool. It is a completely different product '
    'category. The "OpenClaw killer" positioning should be retired.',
    bold=True
)

doc.add_page_break()

# ============================================================
# RAG CLASSIFICATION
# ============================================================
add_heading('Is Claude Brain a RAG System?')

add_para(
    'RAG (Retrieval-Augmented Generation) is the pattern of retrieving relevant data '
    'from a knowledge base, injecting it into an LLM\'s prompt, and generating a response '
    'informed by that context. Claude-brain includes RAG capabilities, but its architecture '
    'extends well beyond what a typical RAG system provides.'
)

add_heading('Where Claude Brain Uses RAG', level=2)
add_bullet(
    'The user-prompt-submit hook runs FTS5 search against the user\'s prompt, '
    'retrieves matching transcripts, and injects them into context automatically.',
    bold_prefix='Automatic context injection: '
)
add_bullet(
    'The session-start hook retrieves previous session summaries and injects them '
    'so Claude has continuity from prior work.',
    bold_prefix='Session continuity: '
)
add_bullet(
    'Claude uses MCP tools (search_transcripts, lookup_fact, lookup_decision, '
    'search_semantic) to retrieve relevant knowledge during a conversation and '
    'incorporate it into responses.',
    bold_prefix='On-demand retrieval via MCP: '
)

add_heading('Where Claude Brain Goes Beyond RAG', level=2)
add_bullet(
    'Facts, decisions, and preferences are stored in dedicated, '
    'queryable tables with categories, project tags, and timestamps. This is a '
    'knowledge base, not a document chunk store.',
    bold_prefix='Structured knowledge: '
)
add_bullet(
    'Weekly email digests and dormant project alerts push information '
    'to the user without any query. RAG is pull; this is push.',
    bold_prefix='Proactive outreach: '
)
add_bullet(
    'Numbered decisions, session protocols, project tracking, and '
    'dependency management. No RAG system provides process management.',
    bold_prefix='Governance: '
)
add_bullet(
    'Every user message and assistant response is preserved and searchable. '
    'Most RAG systems chunk documents, embed them, and discard the originals. '
    'Claude-brain keeps the raw material intact.',
    bold_prefix='Full transcript preservation: '
)
add_bullet(
    'Both FTS5 keyword search and numpy-based cosine similarity '
    'over sentence-transformer embeddings. Most RAG systems use only one approach.',
    bold_prefix='Dual search (keyword + semantic): '
)

add_para('')
add_para(
    'Classification: Claude-brain is a personal AI development governance system '
    'with RAG as one of its retrieval layers. RAG is a component of the architecture, '
    'not the architecture itself. Competitors like Mem0 and claude-mem are closer to '
    'pure RAG systems. Claude-brain adds governance, proactive intelligence, and '
    'structured knowledge on top of the retrieval pattern.',
    bold=True
)

doc.add_page_break()

# ============================================================
# TOTAL RECALL: THE REAL DIFFERENTIATOR
# ============================================================
add_heading('Total Recall: The Real Differentiator')

add_para(
    'The most important architectural difference between claude-brain and every '
    'competitor is not a feature. It is a design philosophy: capture everything, '
    'discard nothing, and let the user decide what matters.'
)

add_heading('The Lossy vs. Lossless Problem', level=2)
add_para(
    'Most AI memory tools use lossy compression. claude-mem captures tool observations '
    'and compresses them with AI. Mem0 extracts "memories" from conversations. '
    'Windsurf auto-generates memory snippets. In every case, an algorithm decides '
    'what is worth keeping and throws away the rest.'
)
add_para(
    'This creates a fundamental limitation: you can only ask questions the extraction '
    'algorithm anticipated. If it did not think a detail was important enough to save, '
    'that detail is gone forever. You cannot go back to a conversation from three '
    'weeks ago and ask "what exact approach did we try before we found the fix?" '
    'because the raw conversation no longer exists.'
)
add_para(
    'Claude-brain takes the opposite approach. Every user message and every assistant '
    'response is preserved in full, searchable via keyword and semantic search, '
    'with structured metadata (session, project, timestamp, speaker). Nothing is '
    'discarded. Nothing is compressed. The raw material is always available.',
    bold=True
)

add_heading('What This Enables', level=2)
add_para(
    'Full transcript capture combined with structured knowledge and natural language '
    'access creates an unlimited query surface. These are not pre-built features. '
    'They are emergent capabilities that arise from the architecture:'
)

add_heading('Lessons Learned (Instant, No Report Required)', level=3)
add_bullet('"What did we try when the MCP server wouldn\'t connect?" — searches '
           'transcripts for that debugging session, returns the exact back-and-forth '
           'including failed approaches and the eventual fix.')
add_bullet('"Show me every time we hit a token limit" — cross-session pattern that '
           'reveals workflow problems invisible in any single conversation.')
add_bullet('"What was the reasoning behind Decision 89?" — pulls the full transcript '
           'context around a specific decision, not just the one-line summary.')

add_heading('Decision Archaeology', level=3)
add_bullet('"When did we decide to drop ChromaDB, and what were the alternatives?" — '
           'retrieves the session, the discussion, and the competing options.')
add_bullet('"Show me every decision that got reversed" — queries decisions table plus '
           'transcript context to surface changed minds and why.')

add_heading('Real-Time Project Status (Just Ask)', level=3)
add_bullet('"What\'s the status of the Johnny Goods project?" — no report to run, no '
           'dashboard to open. MCP tools query facts, recent sessions, and summaries '
           'to give an instant natural-language status update.')
add_bullet('"What did I work on last week across all projects?" — cross-project '
           'session query with summaries, decisions, and open items.')
add_bullet('"Is there anything stuck or blocked?" — dormant project detection plus '
           'session notes with "next step" and "blocker" fields.')

add_heading('Code Pattern Recall', level=3)
add_bullet('"How did I set up the cosine similarity search?" — finds the session '
           'where the code was written and returns the full implementation discussion.')
add_bullet('"What was the MCP server architecture?" — semantic search across all '
           'sessions where MCP was discussed, ordered by relevance.')

add_heading('Cross-Project Intelligence', level=3)
add_bullet('"Find every time I discussed authentication across all projects" — '
           'searches 13,000+ transcripts across 5 projects.')
add_bullet('"Are there patterns in how long my sessions run?" — meta-analysis of '
           'session metadata reveals workflow habits.')

add_heading('The Key Insight', level=2)
add_para(
    'None of the capabilities above were pre-built as features. They emerge from '
    'the combination of full transcript capture, structured knowledge, and flexible '
    'retrieval. This is the fundamental advantage of lossless architecture over lossy '
    'compression: every future question has an answer, including questions the system '
    'designers never anticipated.',
    bold=True
)

add_para('')
add_para(
    'When a competitor extracts "memories" and discards conversations, they are '
    'making a bet that their extraction algorithm captured everything that will ever '
    'matter. That bet is wrong. Development work generates insights that only become '
    'relevant weeks or months later. Debugging approaches that failed on one project '
    'become solutions on another. The developer who can instantly recall any '
    'conversation from any project at any point has a compounding advantage that grows '
    'with every session.',
)

doc.add_page_break()

# ============================================================
# FEATURE COMPARISON
# ============================================================
add_heading('Feature-by-Feature Comparison')

add_para(
    'The following matrix compares claude-brain against its closest competitors '
    'across all relevant capabilities.'
)

add_table(
    ['Feature', 'claude-brain', 'claude-mem', 'Mem0', 'OpenMemory', 'Built-in'],
    [
        ['Full transcript capture', 'Yes', 'No (tool obs)', 'No', 'No', 'No'],
        ['Structured knowledge', 'Yes (3 tables)', 'No', 'Partial', 'No', 'No'],
        ['Cross-project search', 'Yes', 'Basic', 'No', 'No', 'No'],
        ['Semantic search', 'Yes (numpy)', 'Yes (Chroma)', 'Yes', 'Yes', 'No'],
        ['FTS5 keyword search', 'Yes', 'Yes', 'No', 'No', 'No'],
        ['LLM session summaries', 'Yes', 'Yes', 'No', 'No', 'No'],
        ['Slash commands', 'Yes (6)', 'Yes (4)', 'No', 'No', 'No'],
        ['MCP tools', 'Yes (10)', 'Yes', 'Yes', 'Yes', 'No'],
        ['Proactive email digests', 'Yes', 'No', 'No', 'No', 'No'],
        ['Decision tracking', 'Yes', 'No', 'No', 'No', 'No'],
        ['Local / private', 'Yes', 'Yes', 'Cloud', 'Yes', 'Yes'],
        ['Zero cost', 'Yes', 'Yes', 'No ($19/mo+)', 'Yes', 'Yes'],
        ['Background service', 'No (on demand)', 'Yes (port 37777)', 'Yes', 'Yes', 'No'],
    ]
)

doc.add_paragraph()
add_para(
    'No single competitor offers the full combination of features that claude-brain provides. '
    'The closest competitor (claude-mem) is missing full transcript capture, structured '
    'knowledge, proactive features, and governance.',
    bold=True
)

doc.add_page_break()

# ============================================================
# ACQUISITION LANDSCAPE
# ============================================================
add_heading('Acquisition Landscape')

add_heading('Recent AI / Developer Tool Acquisitions', level=2)

add_table(
    ['Acquirer', 'Target', 'Price', 'Relevance'],
    [
        ['Anthropic', 'Bun (JS runtime)', '~$100-200M', 'First acquisition. Bun powers Claude Code.'],
        ['OpenAI', 'OpenClaw creator', 'Undisclosed', 'Acqui-hire for distribution (287K stars).'],
        ['OpenAI', 'Windsurf / Codeium', '$3B (failed)', 'Collapsed over Microsoft IP rights.'],
        ['Google', 'Windsurf leadership', '$2.4B', 'Talent + license deal for 40 senior staff.'],
        ['Cognition', 'Windsurf product', '~$250M', 'Product + 210 employees + $82M ARR.'],
        ['IBM', 'HashiCorp', '$6.4B', 'Largest OSS dev tool acquisition. ~40K stars.'],
        ['Cursor', 'Graphite', '>$290M', 'Code review / stacked PRs.'],
        ['OpenAI', 'Statsig', '$1.1B', 'Product experimentation platform.'],
    ]
)

doc.add_paragraph()

add_heading('Star Count to Acquisition Value', level=2)

add_table(
    ['Stars', 'Example', 'Outcome'],
    [
        ['10K+', 'Enterprise signal', 'VCs pay attention'],
        ['40K+', 'HashiCorp / Terraform', '$6.4B acquisition by IBM'],
        ['75K+', 'Bun', '~$100-200M acquisition by Anthropic'],
        ['180K+', 'OpenClaw', 'Acqui-hire by OpenAI, multiple Big Tech bidders'],
    ]
)

doc.add_paragraph()

add_heading('AI Coding Tool Valuations', level=2)

add_table(
    ['Company', 'Valuation', 'ARR', 'Status'],
    [
        ['Cursor', '$29.3B', '$2B+ (doubling quarterly)', 'Independent, acquiring others'],
        ['Replit', '$9B', '$150M+', 'Independent'],
        ['Windsurf', 'Was $3B', '$82M at split', 'Split: Google + Cognition'],
    ]
)

doc.add_paragraph()

add_heading('What Acquirers Want', level=2)
add_bullet(
    'Stars = developer mindshare. OpenClaw analyst: "They bought a distribution channel."',
    bold_prefix='Distribution / adoption: '
)
add_bullet(
    'Anthropic bought Bun because Claude Code literally runs on it.',
    bold_prefix='Strategic dependency: '
)
add_bullet(
    'Windsurf\'s $82M ARR doubling quarterly made it a $3B target.',
    bold_prefix='Revenue trajectory: '
)
add_bullet(
    'Architecture competitors can\'t easily replicate.',
    bold_prefix='Technical moat: '
)
add_bullet(
    'The "reverse acqui-hire" is the dominant deal structure. $40B+ deployed this way 2024-2026.',
    bold_prefix='Talent: '
)

doc.add_page_break()

# ============================================================
# HONEST ASSESSMENT
# ============================================================
add_heading('Honest Assessment')

add_heading('Strengths', level=2)

add_bullet(
    '51 of 52 steps complete, tested, working, clean. '
    '116 sessions, 13,377 transcripts, semantic search, structured facts and decisions. '
    'This is properly engineered, not a weekend hack.',
    bold_prefix='The architecture is real. '
)
add_bullet(
    'Full transcripts + structured governance + semantic search + proactive outreach + '
    'cross-project isolation + local-first + zero cost. No single competitor offers this.',
    bold_prefix='The combination is genuinely unique. '
)
add_bullet(
    'Every other tool is passive ("store and retrieve"). '
    'Nobody tracks decisions, session protocols, or project state. '
    'Nobody sends proactive email digests.',
    bold_prefix='The governance angle is a real gap in the market. '
)
add_bullet(
    'claude-mem got 33.5K stars proving developers are desperate for Claude Code memory. '
    'The built-in 200-line cap is inadequate.',
    bold_prefix='Demand is proven. '
)
add_bullet(
    'Anthropic bought infrastructure their product depends on. '
    'If claude-brain becomes what power users depend on, the precedent exists.',
    bold_prefix='The Bun acquisition precedent is real. '
)

add_heading('Weaknesses', level=2)

add_bullet(
    'Most developers don\'t track numbered decisions or write session protocols. '
    'The governance features that make this unique may be the features most users skip.',
    bold_prefix='Unique does not mean universally wanted. '
)
add_bullet(
    'claude-mem has 33.5K stars and a head start. Mem0 has $24M and a team. '
    'Getting from 0 to 5K stars as a solo developer against funded competition is hard.',
    bold_prefix='Late entry against funded teams. '
)
add_bullet(
    'If Anthropic bumps MEMORY.md from 200 to 10,000 lines and adds semantic search, '
    '80% of the value proposition is at risk. They haven\'t in 2 years, but could.',
    bold_prefix='Platform risk. '
)
add_bullet(
    'Anthropic bought Bun at 75K stars. OpenAI acqui-hired OpenClaw\'s creator at 287K stars. '
    'Those are massive distribution numbers. The bar is high.',
    bold_prefix='Acquisition requires massive scale. '
)
add_bullet(
    'Hooks + MCP registration + config + Python dependencies + model download. '
    'claude-mem installs from a plugin marketplace in one command.',
    bold_prefix='Setup complexity vs. competitors. '
)

doc.add_page_break()

# ============================================================
# STRATEGY: OPEN SOURCE + TARGETED POSITIONING
# ============================================================
add_heading('Strategy: Open Source with Targeted Positioning')

add_para(
    'After reviewing the competitive landscape, acquisition precedents, and honest '
    'assessment of strengths and weaknesses, the decision is to go public as an '
    'open-source Claude Code power tool with a targeted outreach strategy aimed '
    'at the Anthropic / Claude Code team and the power user community.',
    bold=True
)

add_heading('Why Open Source Is the Right Move', level=2)
add_para(
    'This is not "post it and hope people find it." This is a targeted campaign '
    'to put a production-quality solution in front of the people who matter most: '
    'the Claude Code team at Anthropic and the power user community they listen to.'
)
add_bullet(
    'Boris Cherny (Claude Code creator) and the Anthropic developer experience '
    'team are actively watching what the community builds. They have seen claude-mem '
    'get 33.5K stars. They know their built-in memory (200-line cap, no search) is '
    'inadequate. A properly architected solution that power users depend on gets noticed.',
    bold_prefix='The Anthropic channel is direct. '
)
add_bullet(
    'claude-mem at 33.5K stars proves developers are actively searching for '
    'Claude Code memory solutions. The demand exists and is growing.',
    bold_prefix='Demand is proven. '
)
add_bullet(
    'The pitch to Anthropic\'s team is not "use my tool." It is: "here is a working '
    'architecture for the memory problem your users are asking about. Open source, '
    'MIT licensed. Take it." If they adopt ideas from it, we still win (recognition, '
    'credibility, conversation). If they want to integrate or acquire it, we win bigger.',
    bold_prefix='The pitch writes itself. '
)
add_bullet(
    'No other tool combines full transcript capture, structured governance, '
    'semantic search, proactive email digests, cross-project isolation, and '
    'local-first architecture. This combination is genuinely unique.',
    bold_prefix='The product is differentiated. '
)
add_bullet(
    'Time spent on polish (hook conversion, README, Mac test). '
    'If it does not gain traction, the tool still works for personal use. '
    'The downside is limited. The upside is real.',
    bold_prefix='Risk is low. '
)

add_heading('Positioning', level=2)
add_para(
    'claude-brain is the most complete memory and governance system for Claude Code '
    'power users. It captures every conversation, tracks every decision, searches '
    'across all projects, and proactively reports on your development workflow. '
    'Local, private, and zero cost.'
)
add_para('')
add_para(
    'Tagline: "claude-mem records your tool usage. claude-brain remembers everything."',
    italic=True
)
add_para('')
add_para(
    'The key message: most AI memory tools use lossy compression, deciding what to '
    'remember and discarding the rest. claude-brain keeps everything. Full transcripts, '
    'structured decisions, semantic search, proactive digests. Every past conversation '
    'is instantly queryable. Lessons learned, project status, decision history — '
    'just ask. No reports to run, no dashboards to open.',
)

add_heading('Target Outcomes', level=2)
add_table(
    ['Outcome', 'Likelihood', 'Value'],
    [
        ['1-5K stars, community recognition, portfolio piece', 'High', 'Career leverage, community contributions improve the tool'],
        ['10K+ stars, power user adoption', 'Medium', 'Anthropic conversation becomes likely'],
        ['Anthropic adopts architecture ideas', 'Medium', 'Recognition, credibility, possible hire'],
        ['Anthropic integrates or acquires', 'Low (but real)', 'Bun precedent: buy what Claude Code depends on'],
    ]
)

doc.add_page_break()

# ============================================================
# OUTREACH PLAN
# ============================================================
add_heading('Outreach Plan')

add_para(
    'This is not spray-and-pray marketing. This is targeted outreach to the specific '
    'communities and individuals who will care most about this tool.'
)

add_heading('Tier 1: Anthropic / Claude Code Team (Highest Priority)', level=2)
add_bullet(
    'Boris Cherny is the creator of Claude Code and active on X/Twitter and in '
    'the Claude Code community. A concise demo showing the architecture and what it '
    'solves — tagged directly — puts it on his radar.',
    bold_prefix='Boris Cherny (Claude Code creator): '
)
add_bullet(
    'The Claude Code channel on Anthropic\'s Discord is where power users gather '
    'and the team monitors feedback. A well-crafted post showing the tool in action '
    'gets seen by both users and Anthropic engineers.',
    bold_prefix='Anthropic Discord (Claude Code channel): '
)
add_bullet(
    'Users actively request better memory in Claude Code\'s GitHub issues and '
    'discussions. Linking claude-brain as a solution to existing issues puts it '
    'in the right context.',
    bold_prefix='Claude Code GitHub Discussions: '
)

add_heading('Tier 2: Claude Code Power Users', level=2)
add_bullet(
    '35K+ subscribers, heavily Claude Code focused. Posts showing '
    'real workflow improvements with before/after examples perform well here.',
    bold_prefix='r/ClaudeAI: '
)
add_bullet(
    'Developer audience, drives GitHub stars, reaches '
    'people who build tools. "Show HN" format with a clear problem/solution.',
    bold_prefix='Hacker News (Show HN): '
)
add_bullet(
    'The AI-assisted development community on X. Tag relevant voices: '
    'Claude Code users with large followings, AI dev tool reviewers, etc.',
    bold_prefix='X/Twitter (AI dev community): '
)

add_heading('Tier 3: Broader Developer Community', level=2)
add_bullet(
    'Technical post explaining the architecture and the lossy vs. lossless '
    'insight. Establishes thought leadership.',
    bold_prefix='Dev.to / Medium technical blog post: '
)
add_bullet(
    'If Claude Code supports plugin listings, get listed. '
    'This is where users discover tools.',
    bold_prefix='Plugin marketplace (if available): '
)
add_bullet(
    'Awesome-claude-code lists, AI tool directories, '
    'developer tool newsletters.',
    bold_prefix='Curated lists and directories: '
)

doc.add_page_break()

# ============================================================
# PRE-LAUNCH WORK ITEMS (STEP BY STEP)
# ============================================================
add_heading('Pre-Launch Work Items')

add_para(
    'The following items must be completed before going public. Each step is '
    'designed to be completable in a single session. Work them in order — '
    'each step builds on the previous one.',
    bold=True
)

add_heading('Step 1: Convert Hooks from Bash to Python', level=2)
add_para(
    'All four hooks (session-start, user-prompt-submit, stop, session-end) and '
    'brain_sync.sh must be converted from bash to pure Python. This is required for '
    'cross-platform support (Windows) and signals professional quality.'
)
add_table(
    ['Item', 'Details'],
    [
        ['What', 'Rewrite 4 bash hooks + brain_sync.sh as Python scripts'],
        ['Why', 'Bash hooks = Linux-only. Python = Linux + Mac + Windows. Cross-platform is table stakes.'],
        ['Effort', '~1 session. All bash scripts are thin wrappers around Python already written.'],
        ['Key changes', 'sed/awk -> os.path, ls -t -> glob+sorted, cat /dev/null -> sys.stdin.read()'],
        ['Test', 'Run full session lifecycle (start, prompt, stop, end) with Python hooks on Linux'],
        ['Done when', 'Zero .sh files in hooks/. All hooks registered as .py in settings.json.'],
    ]
)

add_heading('Step 2: Mac Beta Test', level=2)
add_para(
    'A friend (Mac user) clones the repo, runs brain-setup.py, and validates '
    'the setup experience and HOW_TO guide. This proves cross-platform readiness '
    'and catches documentation gaps.'
)
add_table(
    ['Item', 'Details'],
    [
        ['What', 'Friend clones repo on Mac, runs setup, ingests data, tests slash commands'],
        ['Why', 'Cannot claim Mac support without testing on Mac. Documentation gaps are invisible to the author.'],
        ['Effort', '~1 hour of friend\'s time. May require 1 session to fix issues found.'],
        ['Key risks', 'Python path differences, missing dependencies (pip vs pip3), sed differences'],
        ['Test', 'brain-setup.py completes, /brain-status works, hook cycle works, search returns results'],
        ['Done when', 'Friend confirms working on Mac. Issues fixed. Can claim "Linux + macOS supported."'],
    ]
)

add_heading('Step 3: README Rewrite', level=2)
add_para(
    'The current README is technically accurate but reads like internal documentation. '
    'For public launch, it must sell the pain (context loss) and the solution (total recall) '
    'in the first 10 lines. The "Total Recall" and "lossy vs. lossless" framing from this '
    'document should drive the README narrative.'
)
add_table(
    ['Item', 'Details'],
    [
        ['What', 'Rewrite README.md for public audience. Lead with pain, show solution.'],
        ['Why', 'The README is the landing page. Developers decide in 10 seconds whether to star or leave.'],
        ['Effort', '~1 session.'],
        ['Structure', '1) Hook (the problem), 2) What it does (3 bullets), 3) Demo GIF or terminal output, '
         '4) Quick start, 5) How it works, 6) Feature comparison table, 7) Why not claude-mem'],
        ['Key message', '"Most AI memory tools decide what to remember. claude-brain remembers everything."'],
        ['Done when', 'README tells a compelling story in under 60 seconds of reading.'],
    ]
)

add_heading('Step 4: Security and Code Review', level=2)
add_para(
    'Final review before going public. Ensure no personal data, credentials, or '
    'sensitive information in any tracked file. Review .gitignore coverage.'
)
add_table(
    ['Item', 'Details'],
    [
        ['What', 'Scan all tracked files for personal data. Verify .gitignore. Review config.yaml.example.'],
        ['Why', 'One leaked credential or personal detail in a public repo is a permanent mistake.'],
        ['Effort', '~30 minutes. Security audit was done in Step 7.6 but must be re-verified after changes.'],
        ['Check', 'No email addresses, API keys, file paths with usernames, personal names in code/fixtures'],
        ['Done when', 'Every tracked file is safe for public viewing. .gitignore covers all sensitive paths.'],
    ]
)

add_heading('Step 5: Go Public', level=2)
add_para(
    'Switch the GitHub repo from private to public. This is one command but is '
    'irreversible in terms of public visibility. All previous steps must be complete.'
)
add_table(
    ['Item', 'Details'],
    [
        ['What', 'gh repo edit mikeadolan/claude-brain --visibility public'],
        ['Why', 'Cannot get stars, forks, or community while private.'],
        ['Effort', 'One command.'],
        ['Prerequisites', 'Steps 1-4 complete. Security audit passed. README compelling.'],
        ['Done when', 'Repo is publicly accessible. README displays on GitHub.'],
    ]
)

add_heading('Step 6: Launch Posts', level=2)
add_para(
    'Targeted posts to the communities identified in the Outreach Plan. These should '
    'be crafted carefully — each community has different norms and expectations.'
)
add_table(
    ['Item', 'Details'],
    [
        ['What', 'Write and publish posts to Anthropic Discord, r/ClaudeAI, Show HN, X/Twitter'],
        ['Why', 'Stars come from visibility. Visibility comes from the right posts in the right places.'],
        ['Effort', '~1 session to write posts. Stagger over 1-2 weeks for sustained visibility.'],
        ['Discord post', 'Short, focused on what it solves. Link to repo. Ask for feedback.'],
        ['Reddit post', 'Show the workflow. Before/after. Real terminal output or GIF.'],
        ['Show HN', 'Technical angle: the lossy vs. lossless insight. Architecture overview. Link to repo.'],
        ['X/Twitter', 'Tag Boris Cherny. Concise thread: problem, solution, link. No hype.'],
        ['Done when', 'Posts published. Initial reactions monitored. Issues responded to within 24 hours.'],
    ]
)

add_heading('Step 7: Post-Launch (Ongoing)', level=2)
add_para(
    'After launch, the priority shifts to community engagement and iteration.'
)
add_bullet('Respond to GitHub issues within 24 hours')
add_bullet('Accept and review community PRs')
add_bullet('Iterate on feedback — the community will tell you what matters')
add_bullet('Write a CONTRIBUTING.md guide for contributors')
add_bullet('Track star growth: 100 first week, 1K in 90 days is the target')
add_bullet('At month 3: evaluate traction and decide on broader expansion (Option B: universal memory layer)')

doc.add_page_break()

# ============================================================
# DECISION SUMMARY
# ============================================================
add_heading('Decision Summary')

add_para(
    'Go public with open-source Claude Code Power Tool positioning. '
    'Targeted outreach to Boris Cherny, the Claude Code team, and the power user '
    'community. Not spray-and-pray — strategic positioning as the production-quality '
    'answer to Claude Code\'s biggest gap.',
    bold=True
)

add_para('')

add_table(
    ['', 'Decision'],
    [
        ['Strategy', 'Open source, targeted outreach to Anthropic/Claude Code team and power users'],
        ['Positioning', 'Most complete memory + governance system for Claude Code. Total recall, not lossy compression.'],
        ['Tagline', '"claude-mem records your tool usage. claude-brain remembers everything."'],
        ['Key message', '"Most AI memory tools decide what to remember. claude-brain remembers everything."'],
        ['Pre-launch steps', '7 steps: hooks to Python, Mac test, README, security review, go public, launch posts, community engagement'],
        ['Star target', '100 first week, 1K in 90 days'],
        ['Risk', 'Low. Tool works for personal use regardless. Time investment is polish, not new development.'],
        ['Best case', 'Power user adoption forces Anthropic conversation (Bun precedent)'],
    ]
)

add_para('')
add_para(
    'This document was prepared on March 9, 2026. The competitive landscape is '
    'moving fast. Star counts, funding, and product features should be re-verified '
    'before any public launch materials reference them.',
    italic=True
)

# ============================================================
# SOURCES
# ============================================================
doc.add_page_break()
add_heading('Sources')
add_para('Research conducted March 9, 2026 via comprehensive web search across multiple domains.', italic=True)
add_para('')

sources = [
    'Anthropic Acquires Bun — anthropic.com/news',
    'OpenClaw GitHub Repository — github.com/openclaw/openclaw',
    'OpenClaw Wikipedia — en.wikipedia.org/wiki/OpenClaw',
    'OpenClaw Security Crisis — reco.ai/blog',
    'OpenClaw Creator Joins OpenAI — TechCrunch, Feb 2026',
    'claude-mem GitHub — github.com/thedotmack/claude-mem',
    'claude-mem Official Docs — docs.claude-mem.ai',
    'Mem0 GitHub — github.com/mem0ai/mem0',
    'Mem0 Series A — TechCrunch, Oct 2025',
    'Letta / MemGPT GitHub — github.com/letta-ai/letta',
    'Letta $10M Seed — PR Newswire',
    'Graphiti / Zep GitHub — github.com/getzep/graphiti',
    'Khoj GitHub — github.com/khoj-ai/khoj',
    'Cognee $7.5M Seed — cognee.ai/blog',
    'Supermemory GitHub — github.com/supermemoryai/supermemory',
    'Supermemory Funding — TechCrunch, Oct 2025',
    'OpenMemory GitHub — github.com/CaviraOSS/OpenMemory',
    'MemOS GitHub — github.com/MemTensor/MemOS',
    'LangMem GitHub — github.com/langchain-ai/langmem',
    'Claude Code Memory Docs — code.claude.com/docs/en/memory',
    'Claude Code Auto Memory Internals — giuseppegurgone.com',
    'Cursor Rules Documentation — cursor.com/docs/context/rules',
    'Windsurf Cascade Memories — docs.windsurf.com',
    'Copilot Memory — GitHub Blog Changelog, Mar 2026',
    'Cline Memory Bank — docs.cline.bot/features/memory-bank',
    'Aider Conventions — aider.chat/docs/usage/conventions.html',
    'IBM Closes HashiCorp — TechCrunch, Feb 2025',
    'Cursor $2.3B Series D — CNBC, Nov 2025',
    'Cursor $2B+ ARR — TechCrunch, Mar 2026',
    'Windsurf Acquisition Saga — DeepLearning.ai, Fortune, VentureBeat',
    'Replit $9B Valuation — Bloomberg, Jan 2026',
    'State of Agent Memory 2026 — blog.virenmohindra.me',
    'AI Coding Tools Compared 2026 — tldl.io',
]
for s in sources:
    add_bullet(s)

# Save
doc.save(OUTPUT)
print(f"Saved to {OUTPUT}")
